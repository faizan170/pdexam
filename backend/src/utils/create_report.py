from docx import Document

from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import os
from docx.shared import Inches, Cm
from uuid import uuid4

def set_cell_margins(cell, **kwargs):
    """
    cell:  actual cell instance you want to modify
    usage:
        set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in ["top", "start", "bottom", "end"]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)


def create_report(report_data):
    print(report_data)
    # create document
    doc = Document()


    table = doc.add_table(rows=1, cols=1, style="Medium Shading 2 Accent 1")
    table.rows[0].cells[0].text = "Report"
    doc.add_paragraph('')


    # Basic report data
    table = doc.add_table(rows=3, cols=2, style="Table Grid")
    for i, key in enumerate(['User Code', 'Date', 'Time']):
    
        cell = table.rows[i].cells
        set_cell_margins(cell[0], top=100, start=100, bottom=100, end=50)
        cell[0].text = key
        cell[1].text = report_data['fields'].get(key)

    for field_data in report_data['fields']['questions']:
        row = table.add_row()
        cell = row.cells
        set_cell_margins(cell[0], top=100, start=100, bottom=100, end=50)
        cell[0].text = field_data['title']
        cell[1].text = field_data['value']

    doc.add_paragraph('')
    
    for i, row_data in enumerate(report_data['spectograms']):

        table = doc.add_table(rows=3, cols=1, style="Table Grid")
        table.rows[0].cells[0].text = f"Test {i+1}: {row_data['title']}"
        print(row_data)
        # add image to table
        if row_data['spectogram']['right'] != "":
            paragraph = table.rows[1].cells[0].paragraphs[0]
            paragraph.add_run(text=f"Right {row_data['type']} Spectogram (10 seconds)\n")
            run = paragraph.add_run()
            run.add_picture(row_data['spectogram']['right'], width=Inches(5.7), height=Inches(1))
            os.remove(row_data['spectogram']['right'])

        # add image to table
        if row_data['spectogram']['left'] != "":
            paragraph = table.rows[2].cells[0].paragraphs[0]
            paragraph.add_run(text=f"Left {row_data['type']} Spectogram (10 seconds)\n")
            run = paragraph.add_run()
            run.add_picture(row_data['spectogram']['left'], width=Inches(5.7), height=Inches(1))
            os.remove(row_data['spectogram']['left'])
        
        if i in [1, 4, 7]:
            doc.add_page_break()
        else:
            doc.add_paragraph('')
    


    report_path = "static/temp/" + str(uuid4()) + ".docx"
    doc.save(report_path)
    return report_path